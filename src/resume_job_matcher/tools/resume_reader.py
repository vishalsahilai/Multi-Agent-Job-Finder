import re
import io
from pathlib import Path
from typing import Union

#  PDF 
def _read_pdf(file: Union[str, Path, bytes]) -> str:
    """Extract text from a PDF file or bytes object."""
    try:
        from pypdf import PdfReader
    except ImportError:
        raise ImportError("pypdf is required: pip install pypdf")
 
    if isinstance(file, (str, Path)):
        reader = PdfReader(str(file))
    else:
        reader = PdfReader(io.BytesIO(file))
 
    pages = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            pages.append(text)
 
    return "\n".join(pages)

#  DOCX
 
def _read_docx(file: Union[str, Path, bytes]) -> str:
    """Extract text from a DOCX file or bytes object."""
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx is required: pip install python-docx")
 
    if isinstance(file, (str, Path)):
        doc = Document(str(file))
    else:
        doc = Document(io.BytesIO(file))
 
    paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
    return "\n".join(paragraphs)

#  Text Cleaner
 
def _clean_text(text: str) -> str:
    """
    Normalize and clean extracted resume text:
    - Fix encoding artifacts (e.g. â€™ → ')
    - Remove null bytes and control characters
    - Collapse 3+ blank lines into 2
    - Strip trailing whitespace per line
    - Normalize bullet characters to '-'
    """
    if not text:
        return ""
 
    # Remove null bytes and non-printable control chars (keep \n \t)
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", text)
 
    # Fix common UTF-8 mojibake artifacts
    replacements = {
        "\u2019": "'",   # right single quotation mark
        "\u2018": "'",   # left single quotation mark
        "\u201c": '"',   # left double quotation mark
        "\u201d": '"',   # right double quotation mark
        "\u2013": "-",   # en dash
        "\u2014": "-",   # em dash
        "\u2022": "-",   # bullet
        "\u2023": "-",   # triangular bullet
        "\u25cf": "-",   # black circle bullet
        "\u00a0": " ",   # non-breaking space
        "\ufeff": "",    # BOM
    }
    for bad, good in replacements.items():
        text = text.replace(bad, good)
 
    # Normalize bullet-like characters at line start
    text = re.sub(r"^[\*\•\·\◦\▪\▸\►\-–—]\s*", "- ", text, flags=re.MULTILINE)
 
    # Strip trailing whitespace per line
    lines = [line.rstrip() for line in text.split("\n")]
 
    # Collapse more than 2 consecutive blank lines into 2
    cleaned_lines = []
    blank_count = 0
    for line in lines:
        if line.strip() == "":
            blank_count += 1
            if blank_count <= 2:
                cleaned_lines.append("")
        else:
            blank_count = 0
            cleaned_lines.append(line)
 
    return "\n".join(cleaned_lines).strip()


 # Public API 
 
def read_resume(file: Union[str, Path, bytes], filename: str = "") -> dict:

    result = {
        "text": "",
        "format": "",
        "char_count": 0,
        "word_count": 0,
        "error": None,
    }
 
    try:
        # Detect format
        if isinstance(file, (str, Path)):
            suffix = Path(file).suffix.lower()
        else:
            suffix = Path(filename).suffix.lower() if filename else ""
 
        if suffix == ".pdf":
            raw_text = _read_pdf(file)
            result["format"] = "pdf"
        elif suffix in (".docx", ".doc"):
            raw_text = _read_docx(file)
            result["format"] = "docx"
        else:
            result["error"] = f"Unsupported file format: '{suffix}'. Upload a PDF or DOCX."
            return result
 
        if not raw_text.strip():
            result["error"] = "Could not extract any text from the resume. It may be image-based or corrupted."
            return result
 
        cleaned = _clean_text(raw_text)
        result["text"] = cleaned
        result["char_count"] = len(cleaned)
        result["word_count"] = len(cleaned.split())
 
    except Exception as e:
        result["error"] = f"Failed to read resume: {str(e)}"
 
    return result
 